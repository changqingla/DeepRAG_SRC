#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from deepdoc.parser.pdf_parser import PlainParser
from rag.nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec, concat_img, \
    naive_merge_docx, tokenize_chunks_docx
from deepdoc.parser import PdfParser, ExcelParser, DocxParser, HtmlParser, JsonParser, MarkdownParser, TxtParser
from rag.utils import num_tokens_from_string
from PIL import Image
from functools import reduce
from markdown import markdown
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        tbls = self._extract_table_figure(True, zoomin, True, True)
        # self._naive_vertical_merge()
        self._concat_downward()
        # self._filter_forpages()

        logging.info("layouts cost: {}s".format(timer() - first_start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def __call__(self, filename, binary=None):
        # 读取文件内容：支持本地文件路径或二进制流
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        # 从 Markdown 中提取表格和正文部分
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')

        sections = []  # 存储正文段落 (text, "")
        tbls = []      # 存储表格 ((None, html_table), "")

        # -------------------------
        # 文本切分部分
        # -------------------------
        for sec in remainder.split("\n"):
            tokens = num_tokens_from_string(sec)

            if tokens  > 7 * self.chunk_token_num:
                # 对超长段落先按句号类标点切分
                sentences = re.split(r'(?<=[.!?])\s+', sec)
                current_chunk = ""
                current_tokens = 0

                for sentence in sentences:
                    sentence_tokens = num_tokens_from_string(sentence)

                    if sentence_tokens > 4 * self.chunk_token_num:
                        # 如果句子本身太长，则进一步用逗号类标点切分
                        parts = re.split(r'(?<=[,;:])\s+', sentence)
                        for part in parts:
                            part_tokens = num_tokens_from_string(part)
                            if current_tokens + part_tokens > self.chunk_token_num:
                                if current_chunk:
                                    sections.append((current_chunk, ""))
                                current_chunk = part
                                current_tokens = part_tokens
                            else:
                                current_chunk += " " + part if current_chunk else part
                                current_tokens += part_tokens
                    else:
                        # 累加当前句子，如果超出 chunk 限制则新开段
                        if current_tokens + sentence_tokens > self.chunk_token_num:
                            if current_chunk:
                                sections.append((current_chunk, ""))
                            current_chunk = sentence
                            current_tokens = sentence_tokens
                        else:
                            current_chunk += " " + sentence if current_chunk else sentence
                            current_tokens += sentence_tokens

                # 收尾：最后剩余的 chunk 也加入
                if current_chunk:
                    sections.append((current_chunk, ""))

            else:
                # 处理正常段落
                if sec.strip().startswith("#"):
                    # 是标题，单独作为一个 section
                    sections.append((sec, ""))
                elif sections and sections[-1][0].strip().startswith("#"):
                    # 如果上一个是标题，尝试与当前正文合并
                    sec_, _ = sections.pop(-1)
                    combined = sec_ + "\n" + sec
                    combined_tokens = num_tokens_from_string(combined)

                    if combined_tokens > self.chunk_token_num:
                        # 合并后过长，不合并
                        sections.append((sec_, ""))
                        sections.append((sec, ""))
                    else:
                        # 合并后合理，则放在一起
                        sections.append((combined, ""))
                else:
                    # 普通段落，直接添加
                    sections.append((sec, ""))

        # -------------------------
        # 表格切分部分
        # -------------------------
        for table in tables:
            html_table = markdown(table, extensions=['markdown.extensions.tables'])
            table_tokens = num_tokens_from_string(html_table)
            
            # 提取表头和表体
            header_match = re.search(r'<thead>(.*?)</thead>|<tr>(.*?</th>.*?)</tr>', html_table, re.DOTALL)
            header = header_match.group(0) if header_match else ""
            header_tokens = num_tokens_from_string(header) if header else 0
            
            if table_tokens > 7 * self.chunk_token_num:
                # 分割表格时保留表头结构
                table_body = html_table.replace(header, "") if header else html_table
                rows = re.findall(r'<tr>.*?</tr>', table_body, re.DOTALL)
                
                # 计算每个分块能容纳的行数
                max_tokens_per_chunk = 7 * self.chunk_token_num - header_tokens
                current_chunk_rows = []
                current_chunk_tokens = header_tokens + num_tokens_from_string("<table></table>")
                
                for row in rows:
                    row_tokens = num_tokens_from_string(row)
                    if current_chunk_tokens + row_tokens > max_tokens_per_chunk:
                        # 保存当前分块
                        tbls.append((
                            (None, f"<table>{header}{''.join(current_chunk_rows)}</table>"), 
                            ""
                        ))
                        # 开始新分块
                        current_chunk_rows = [row]
                        current_chunk_tokens = header_tokens + row_tokens + num_tokens_from_string("<table></table>")
                    else:
                        current_chunk_rows.append(row)
                        current_chunk_tokens += row_tokens
                
                # 添加最后一个分块
                if current_chunk_rows:
                    tbls.append((
                        (None, f"<table>{header}{''.join(current_chunk_rows)}</table>"), 
                        ""
                    ))
            else:
                tbls.append(((None, html_table), ""))

        # 返回正文和表格两个部分
        return sections, tbls 
 

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, excel, txt.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """

    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Docx()(filename, binary)
        res = tokenize_table(tables, doc, is_english)  # just for table

        callback(0.8, "Finish parsing.")
        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_docx(chunks, doc, is_english, images))
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if parser_config.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page,
                                      callback=callback)
        res = tokenize_table(tables, doc, is_english)

    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = ExcelParser()
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Markdown(int(parser_config.get("chunk_token_num", 128)))(filename, binary)
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    chunks = naive_merge(
        sections, int(parser_config.get(
            "chunk_token_num", 128)), parser_config.get(
            "delimiter", "\n!?。；！？"))
    if kwargs.get("section_only", False):
        return chunks

    res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    logging.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)


